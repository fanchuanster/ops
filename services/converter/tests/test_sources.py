"""Tests for the input adapters.

Every accepted input converges on the DOCX master, so the test that
matters most is the round trip: a Document written to a master and read
back must be the same document. That property is what lets the format
stage generate EPUB and PDF from the *approved* master (CLAUDE.md
section 5) rather than from the scan.
"""

from __future__ import annotations

import sys
import zipfile
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.docx.builder import build_docx  # noqa: E402
from app.models import Block, BlockKind, Document  # noqa: E402
from app.pipeline.render import classify_pages  # noqa: E402
from app.sources import SourceKind, UnsupportedSource, load_source  # noqa: E402
from app.sources.detect import detect  # noqa: E402
from app.sources.docx_in import read_docx  # noqa: E402
from app.sources.pdf_in import read_pdf  # noqa: E402
from app.sources.text import read_text  # noqa: E402


def write_pdf(path: Path, pages: list[str]) -> Path:
    """A PDF whose pages are each `text`, `image` or `blank`.

    Real books mix them — a scan with a born-digital colophon, a digital
    book with plates and blank versos — and every rule about what gets
    OCR'd is a rule about one page, so the fixture has to be able to say
    what each page is.
    """
    import pymupdf

    doc = pymupdf.open()
    for index, kind in enumerate(pages):
        page = doc.new_page()
        if kind == "text":
            page.insert_text((72, 120), f"Page {index}.", fontsize=11)
        elif kind == "image":
            # No text layer, something to look at: a scanned leaf.
            pixmap = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, 16, 16), False)
            pixmap.clear_with(255)
            page.insert_image(pymupdf.Rect(72, 72, 172, 172), pixmap=pixmap)
    doc.save(str(path))
    doc.close()
    return path


# --------------------------------------------------------------------------
# The DOCX master round trip
# --------------------------------------------------------------------------


def sample_document() -> Document:
    return Document(
        title="靜夜思選",
        author="李白",
        blocks=[
            Block(kind=BlockKind.CHAPTER, lines=["第一章"], page=0),
            Block(kind=BlockKind.SECTION, lines=["一、月與鄉思"], page=1),
            Block(kind=BlockKind.MARKER, lines=["（一）"], page=1),
            Block(
                kind=BlockKind.VERSE,
                lines=["床前明月光", "疑是地上霜", "舉頭望明月", "低頭思故鄉"],
                page=1,
            ),
            Block(kind=BlockKind.ATTRIBUTION, lines=["——唐·李白《靜夜思》"], page=1),
            Block(
                kind=BlockKind.BODY,
                lines=["這是一段散文，講述詩人的心境。"],
                page=1,
                starts_paragraph=True,
            ),
            Block(kind=BlockKind.FOOTNOTE, lines=["*見第七十一頁。"], page=1),
        ],
    )


def test_a_master_reads_back_as_the_document_that_wrote_it(tmp_path):
    original = sample_document()
    build_docx(original, tmp_path / "master.docx")
    restored = read_docx(tmp_path / "master.docx")

    assert restored.title == original.title
    assert restored.author == original.author
    assert [(b.kind, b.lines) for b in restored.blocks] == [
        (b.kind, b.lines) for b in original.blocks
    ]


def test_a_poem_survives_the_round_trip_as_one_block(tmp_path):
    # The failure this guards against is silent: verse flattened into
    # prose still produces a readable book, just not the poem.
    build_docx(sample_document(), tmp_path / "master.docx")
    restored = read_docx(tmp_path / "master.docx")

    verses = [b for b in restored.blocks if b.kind is BlockKind.VERSE]
    assert len(verses) == 1
    assert len(verses[0].lines) == 4


def test_a_source_reference_reattaches_to_its_block(tmp_path):
    doc = Document(
        title="t",
        blocks=[
            Block(
                kind=BlockKind.ATTRIBUTION,
                lines=["——唐·李白"],
                page=1,
                source_ref="（見第 71 頁）",
            )
        ],
    )
    build_docx(doc, tmp_path / "m.docx")
    restored = read_docx(tmp_path / "m.docx")

    assert restored.blocks[0].source_ref == "（見第 71 頁）"


def test_two_prose_paragraphs_stay_two_paragraphs(tmp_path):
    doc = Document(
        title="t",
        blocks=[
            Block(kind=BlockKind.BODY, lines=["第一段。"], page=0, starts_paragraph=True),
            Block(kind=BlockKind.BODY, lines=["第二段。"], page=0, starts_paragraph=True),
        ],
    )
    build_docx(doc, tmp_path / "m.docx")
    restored = read_docx(tmp_path / "m.docx")

    assert [b.lines for b in restored.blocks] == [["第一段。"], ["第二段。"]]


def test_a_section_head_does_not_come_back_as_a_chapter(tmp_path):
    """The failure this guards against is a corrected master returning as
    a different book: every subheading starting a new page and a new EPUB
    document, and a table of contents claiming chapters the book has not
    got. Silent, and only visible once someone opens the result."""
    build_docx(sample_document(), tmp_path / "master.docx")
    restored = read_docx(tmp_path / "master.docx")

    kinds = [b.kind for b in restored.blocks]
    assert kinds.count(BlockKind.SECTION) == 1
    assert kinds.count(BlockKind.CHAPTER) == 1


def test_a_foreign_docx_falls_back_to_headings_and_prose(tmp_path):
    from docx import Document as DocxDocument

    docx = DocxDocument()
    docx.add_paragraph("第一章", style="Heading 1")
    docx.add_paragraph("第一節", style="Heading 2")
    docx.add_paragraph("尋常的段落。")
    docx.save(str(tmp_path / "foreign.docx"))

    restored = read_docx(tmp_path / "foreign.docx", title="外來文件")

    assert [b.kind for b in restored.blocks] == [
        BlockKind.CHAPTER,
        BlockKind.SECTION,
        BlockKind.BODY,
    ]
    assert restored.title == "外來文件"


def test_a_deeper_heading_is_a_section_rather_than_body(tmp_path):
    # An editor working in Word may nest further than the pipeline's two
    # levels. Losing the heading entirely would be worse than flattening
    # it: the text would come back as a sentence in the middle of a
    # paragraph run.
    from docx import Document as DocxDocument

    docx = DocxDocument()
    docx.add_paragraph("更深的標題", style="Heading 4")
    docx.save(str(tmp_path / "deep.docx"))

    assert read_docx(tmp_path / "deep.docx").blocks[0].kind is BlockKind.SECTION


def test_read_text_is_exact_about_confidence(tmp_path):
    # Text read from a document is not a guess, and must not be filed
    # among the lines a reviewer needs to check.
    from docx import Document as DocxDocument

    docx = DocxDocument()
    docx.add_paragraph("內容")
    docx.save(str(tmp_path / "d.docx"))

    assert all(b.confidence == 1.0 for b in read_docx(tmp_path / "d.docx").blocks)


# --------------------------------------------------------------------------
# Plain text
# --------------------------------------------------------------------------


def test_blank_lines_separate_paragraphs():
    doc = read_text(content="第一段的第一行\n第一段的第二行\n\n第二段。", title="t")

    assert len(doc.blocks) == 2
    assert doc.blocks[0].lines == ["第一段的第一行", "第一段的第二行"]


def test_a_markdown_heading_becomes_a_chapter():
    doc = read_text(content="# 第一章\n\n正文。", title="t")

    assert doc.blocks[0].kind is BlockKind.CHAPTER
    assert doc.blocks[0].lines == ["第一章"]


def test_a_lone_poem_marker_is_recognised():
    doc = read_text(content="（十一）\n\n正文。", title="t")

    assert doc.blocks[0].kind is BlockKind.MARKER


def test_short_lines_are_not_guessed_to_be_verse():
    # Classical prose is set in short lines too. Guessing here would
    # shatter a paragraph into fake poetry in a published book.
    doc = read_text(content="子曰\n學而時習之\n不亦說乎", title="t")

    assert len(doc.blocks) == 1
    assert doc.blocks[0].kind is BlockKind.BODY


def test_the_title_falls_back_to_the_filename(tmp_path):
    path = tmp_path / "論語.txt"
    path.write_text("正文。", encoding="utf-8")

    assert read_text(path).title == "論語"


# --------------------------------------------------------------------------
# Detection
# --------------------------------------------------------------------------


def test_a_docx_is_detected_by_content_not_by_name(tmp_path):
    from docx import Document as DocxDocument

    misnamed = tmp_path / "book.pdf"
    docx = DocxDocument()
    docx.add_paragraph("內容")
    docx.save(str(misnamed))

    assert detect(misnamed) is SourceKind.DOCX


def test_a_plain_zip_is_refused_rather_than_read_as_a_document(tmp_path):
    archive = tmp_path / "stuff.zip"
    with zipfile.ZipFile(archive, "w") as z:
        z.writestr("readme.txt", "hello")

    with pytest.raises(UnsupportedSource, match="not a DOCX"):
        detect(archive)


def test_a_utf8_text_file_is_detected(tmp_path):
    path = tmp_path / "book.txt"
    path.write_text("學而時習之，不亦說乎？", encoding="utf-8")

    assert detect(path) is SourceKind.TEXT


def test_a_chinese_text_file_is_not_refused_over_the_sample_boundary(tmp_path):
    """The bug this guards: a real book refused as "not a plain text file".

    The sniff reads the first 8192 bytes, which for a CJK file lands in
    the middle of a 3-byte character two times in three. The old retry
    dropped exactly three bytes and landed mid-character just as often,
    so a perfectly good UTF-8 book was refused outright.
    """
    path = tmp_path / "sunzi.txt"
    # Every character is three bytes, so no multiple of the sample size
    # falls on a boundary: 8192 is not divisible by 3.
    path.write_text("孫子兵法" * 3000, encoding="utf-8")

    assert detect(path) is SourceKind.TEXT


def test_a_mis_encoded_chinese_file_is_refused_rather_than_mangled(tmp_path):
    # GB18030 decoded as UTF-8 does not fail cleanly — it produces
    # mojibake that would flow into a published book.
    path = tmp_path / "big5.txt"
    path.write_bytes("學而時習之，不亦說乎？".encode("gb18030"))

    with pytest.raises(UnsupportedSource):
        detect(path)


def test_a_binary_file_is_refused(tmp_path):
    path = tmp_path / "image.bin"
    path.write_bytes(b"\x89PNG\r\n\x1a\n\x00\x00\x00\x0d")

    with pytest.raises(UnsupportedSource):
        detect(path)


def test_a_missing_file_is_refused(tmp_path):
    with pytest.raises(UnsupportedSource, match="not a file"):
        detect(tmp_path / "nope.txt")


# --------------------------------------------------------------------------
# The dispatcher
# --------------------------------------------------------------------------


def test_load_source_carries_the_author_through(tmp_path):
    path = tmp_path / "book.txt"
    path.write_text("正文。", encoding="utf-8")

    result = load_source(path, title="書", author="作者")

    assert result.kind is SourceKind.TEXT
    assert result.document.author == "作者"
    assert result.document.title == "書"


def test_a_scan_is_refused_and_told_where_scans_go(tmp_path):
    # There is no local OCR. A page with no text layer has to be *read*,
    # and reading pages is Adobe's Export PDF, called from the web
    # application — so this refuses rather than half-reading the book.
    pdf = write_pdf(tmp_path / "scan.pdf", ["image"])

    with pytest.raises(UnsupportedSource, match="Adobe"):
        load_source(pdf)


def test_one_scanned_page_in_a_digital_book_is_not_quietly_dropped(tmp_path):
    """The bug this whole page-by-page split exists to close.

    Detection used to sample the document and ask whether *any* page had
    text. A 400-page scan with one born-digital title page passed, was
    read through the text-layer path, and converted "successfully" into a
    one-page book — every scanned page yielded nothing and vanished. Now
    the scanned page is what decides, and the book is refused rather than
    silently gutted.
    """
    pdf = write_pdf(tmp_path / "mixed.pdf", ["text", "image", "text"])

    assert detect(pdf) is SourceKind.PDF_SCANNED
    with pytest.raises(UnsupportedSource, match="1 of its 3 pages"):
        load_source(pdf)


def test_a_blank_page_does_not_send_a_digital_book_to_ocr(tmp_path):
    # A chapter verso is not a scan. OCR would spend seconds on it to
    # return nothing, and — before per-page classification — its presence
    # would have condemned the whole book to the expensive path.
    pdf = write_pdf(tmp_path / "versos.pdf", ["text", "blank", "text"])

    assert detect(pdf) is SourceKind.PDF_TEXT

    result = load_source(pdf)

    assert "2 of 3 pages read from the text layer, 1 blank" in result.notes[0]


def test_pages_are_classified_one_by_one(tmp_path):
    pdf = write_pdf(tmp_path / "book.pdf", ["text", "image", "blank", "text"])

    sources = classify_pages(pdf)

    assert sources.text == (0, 3)
    assert sources.ocr == (1,)
    assert sources.blank == (2,)
    assert sources.total == 4


def test_a_book_is_never_half_read(tmp_path):
    """The bug the page-by-page classification exists to prevent.

    Detection used to sample the document and ask whether *any* page had
    text. A 400-page scan with one born-digital title page passed, was
    read through the text path, and converted "successfully" into a
    one-page book — every scanned page yielded nothing and vanished.
    """
    pdf = write_pdf(tmp_path / "mostly-scan.pdf", ["text", "image", "image", "image"])

    with pytest.raises(UnsupportedSource, match="3 of its 4 pages"):
        read_pdf(pdf, title="Mixed")


def test_blank_pages_do_not_make_a_book_unreadable(tmp_path):
    # A chapter verso carries neither text nor image. It is skipped, not
    # counted as a page nothing can read.
    pdf = write_pdf(tmp_path / "versos.pdf", ["text", "blank", "text"])

    document, _report, sources = read_pdf(pdf, title="Digital")

    assert sources.blank == (1,)
    assert "Page 0." in "\n".join(b.text for b in document.blocks)


def test_load_source_reads_a_born_digital_pdf_without_ocr(tmp_path):
    import pymupdf

    pdf = tmp_path / "digital.pdf"
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 120), "Hello, reader.", fontsize=11)
    doc.save(str(pdf))
    doc.close()

    result = load_source(pdf, title="Digital")

    assert result.kind is SourceKind.PDF_TEXT
    assert "Hello, reader." in "\n".join(b.text for b in result.document.blocks)


def test_a_born_digital_pdf_is_not_punctuation_normalized(tmp_path):
    # The normalization rules repair OCR mistakes. Exact text has no
    # mistakes to repair, and "repairing" it would edit the author.
    import pymupdf

    pdf = tmp_path / "digital.pdf"
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 120), "see (note) here", fontsize=11)
    doc.save(str(pdf))
    doc.close()

    result = load_source(pdf)
    text = "\n".join(b.text for b in result.document.blocks)

    assert "(note)" in text, "half-width parens must survive exact text"
    assert "（note）" not in text
