"""DOCX → Document.

Two jobs, and the second is the one that matters most.

The obvious one is reading a DOCX a reader uploaded, so it can become a
NobleSee master like any other input.

The load-bearing one is reading back a master this pipeline wrote. The
DOCX master is the source of truth, and reader-facing formats must be
generated from the *approved* master rather than from the scan (CLAUDE.md
section 5). That only means anything if the approved file can be read
back — which is why the builder writes real named styles instead of
direct formatting. Those style names are the structure, and this maps
them back.
"""

from __future__ import annotations

from pathlib import Path

from ..models import Block, BlockKind, Document

# The inverse of `docx/builder.py`'s `_STYLE_FOR`. Kept beside it in
# spirit; the round-trip test in tests/test_sources.py is what keeps them
# honest, because a silent mismatch here would quietly flatten a book's
# verse into prose.
STYLE_TO_KIND = {
    "NobleSee Marker": BlockKind.MARKER,
    "NobleSee Verse": BlockKind.VERSE,
    "NobleSee Attribution": BlockKind.ATTRIBUTION,
    "NobleSee Footnote": BlockKind.FOOTNOTE,
    "NobleSee Body": BlockKind.BODY,
}

# Headings are Word's own built-in styles rather than NobleSee ones, so
# that the master carries a real outline — one an editor can navigate in
# the sidebar, and one a DOCX from anywhere else already speaks.
#
# The *level* has to survive. A section head read back as a chapter would
# start a new page and a new EPUB document at every subheading, and the
# reader would get a table of contents claiming forty chapters where the
# book has six. That is not a cosmetic difference: it is the corrected
# master coming back as a different book from the one the editor
# approved.
HEADING_KINDS = {
    "Title": BlockKind.CHAPTER,
    "Heading 1": BlockKind.CHAPTER,
}


def _heading_kind(style: str) -> BlockKind | None:
    """CHAPTER, SECTION, or None if this style is not a heading at all.

    Anything below Heading 1 is a SECTION rather than a deeper kind of
    its own. The pipeline has two heading levels because that is what the
    OCR side can honestly distinguish (`classifyParagraphs` in
    `apps/web/src/domain/ocr.ts`), and inventing a third here would be a
    level nothing upstream can fill.
    """
    known = HEADING_KINDS.get(style)
    if known is not None:
        return known
    return BlockKind.SECTION if style.startswith("Heading ") else None

REFERENCE_STYLE = "NobleSee Reference"


def read_docx(path: Path, title: str | None = None) -> Document:
    """Read a DOCX into the pipeline's own structure."""
    from docx import Document as DocxDocument

    docx = DocxDocument(str(path))
    props = docx.core_properties

    doc = Document(
        title=title or (props.title or path.stem),
        author=props.author or None,
    )

    for paragraph in docx.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style = paragraph.style.name if paragraph.style is not None else ""

        # The document title is metadata, not a block: re-emitting it
        # would put the title in the body on every round trip.
        if style == "Title" and text == doc.title:
            continue

        # A reference belongs to the block above it — that is how the
        # builder emitted it, and re-attaching it here is what makes the
        # round trip lossless.
        if style == REFERENCE_STYLE:
            if doc.blocks:
                previous = doc.blocks[-1]
                previous.source_ref = (previous.source_ref or "") + text
            continue

        kind = STYLE_TO_KIND.get(style)
        if kind is None:
            kind = _heading_kind(style) or BlockKind.BODY

        # Consecutive verse paragraphs are one poem. Everything else
        # stands alone: two adjacent BODY paragraphs are two paragraphs,
        # and merging them would destroy the author's own break.
        if kind is BlockKind.VERSE and doc.blocks and doc.blocks[-1].kind is BlockKind.VERSE:
            doc.blocks[-1].lines.append(text)
            continue

        doc.blocks.append(
            Block(
                kind=kind,
                lines=[text],
                # A DOCX has no pages until it is laid out, so there is no
                # page number to record. 0 rather than a fabricated one.
                page=0,
                # Text read from a document is exact. Confidence is an OCR
                # concept and does not apply.
                confidence=1.0,
                starts_paragraph=kind is BlockKind.BODY,
            )
        )

    return doc
