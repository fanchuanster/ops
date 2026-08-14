#!/usr/bin/env python3
"""Generate the reader-facing files for NobleSee's seed books.

This is a one-off developer utility, NOT part of the runtime stack. It
stands in for the real AI/OCR conversion pipeline described in CLAUDE.md
sections 7-11 (see docs/ROADMAP.md), using the same tooling that pipeline
is expected to use, so the seed content is reproducible instead of being
a pile of committed binaries nobody can regenerate.

Output lands in content/seed/. Those files are also mirrored into the R2
bucket, which is where apps/web/src/seed/seed.ts points the catalog's
artifact keys.

For each part it produces the DOCX master plus the reader-facing formats
generated *from* that same content: EPUB, and PDF in three font sizes.

    pip install python-docx ebooklib weasyprint pillow
    python3 tools/generate-seed-content.py            # only missing files
    python3 tools/generate-seed-content.py --force    # regenerate all
    python3 tools/generate-seed-content.py --book analects

NOTE ON TEXT FIDELITY: the passages below are transcribed for
demonstration purposes. NobleSee's whole point is faithful
reproduction, so any text destined for real publication must be
proofread against an authoritative source first — see the human-review
step in CLAUDE.md section 7.
"""

import argparse
import os
import sys

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ebooklib import epub
from weasyprint import HTML
from PIL import Image, ImageDraw, ImageFont

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SEED_ROOT = os.path.join(REPO_ROOT, "content", "seed")

CJK_FONT = "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf"
LATIN_FONT = "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf"

PDF_VARIANTS = {"pdf-standard": 12, "pdf-large": 16, "pdf-xl": 22}


# --------------------------------------------------------------------------
# Book specs
# --------------------------------------------------------------------------

ANALECTS = {
    "slug": "analects",
    "title": "The Analects (論語)",
    "cover_zh": "論語",
    "cover_en": "The Analects",
    "author": "Confucius (孔子)",
    "translator": "James Legge (1893)",
    "colophon": (
        "Translation by James Legge, The Chinese Classics, vol. 1 "
        "(Oxford: Clarendon Press, 1893). This translation is in the "
        "public domain."
    ),
    "parts": [
        {
            "slug": "book-1",
            "title": "Book I — 學而 (Xue Er)",
            "sections": [
                {
                    "zh": "子曰：「學而時習之，不亦說乎？有朋自遠方來，不亦樂乎？"
                          "人不知而不慍，不亦君子乎？」",
                    "en": "The Master said, 'Is it not pleasant to learn with a "
                          "constant perseverance and application? Is it not "
                          "delightful to have friends coming from distant "
                          "quarters? Is he not a man of complete virtue, who "
                          "feels no discomposure though men may take no note "
                          "of him?'",
                },
                {
                    "zh": "子曰：「巧言令色，鮮矣仁！」",
                    "en": "The Master said, 'Fine words and an insinuating "
                          "appearance are seldom associated with true virtue.'",
                },
                {
                    "zh": "曾子曰：「吾日三省吾身：為人謀而不忠乎？"
                          "與朋友交而不信乎？傳不習乎？」",
                    "en": "The philosopher Tsang said, 'I daily examine myself "
                          "on three points: — whether, in transacting business "
                          "for others, I may have been not faithful; — whether, "
                          "in intercourse with friends, I may have been not "
                          "sincere; — whether I may have not mastered and "
                          "practised the instructions of my teacher.'",
                },
            ],
        },
        {
            "slug": "book-2",
            "title": "Book II — 為政 (Wei Zheng)",
            "sections": [
                {
                    "zh": "子曰：「為政以德，譬如北辰，居其所而眾星共之。」",
                    "en": "The Master said, 'He who exercises government by "
                          "means of his virtue may be compared to the north "
                          "polar star, which keeps its place and all the stars "
                          "turn towards it.'",
                },
                {
                    "zh": "子曰：「道之以政，齊之以刑，民免而無恥；"
                          "道之以德，齊之以禮，有恥且格。」",
                    "en": "The Master said, 'If the people be led by laws, and "
                          "uniformity sought to be given them by punishments, "
                          "they will try to avoid the punishment, but have no "
                          "sense of shame. If they be led by virtue, and "
                          "uniformity sought to be given them by the rules of "
                          "propriety, they will have the sense of shame, and "
                          "moreover will become good.'",
                },
                {
                    "zh": "子曰：「吾十有五而志于學，三十而立，四十而不惑，"
                          "五十而知天命，六十而耳順，七十而從心所欲，不踰矩。」",
                    "en": "The Master said, 'At fifteen, I had my mind bent on "
                          "learning. At thirty, I stood firm. At forty, I had "
                          "no doubts. At fifty, I knew the decrees of Heaven. "
                          "At sixty, my ear was an obedient organ for the "
                          "reception of truth. At seventy, I could follow what "
                          "my heart desired, without transgressing what was "
                          "right.'",
                },
            ],
        },
        {
            "slug": "book-3",
            "title": "Book III — 八佾 (Ba Yi)",
            "sections": [
                {
                    "zh": "子曰：「人而不仁，如禮何？人而不仁，如樂何？」",
                    "en": "The Master said, 'If a man be without the virtues "
                          "proper to humanity, what has he to do with the rites "
                          "of propriety? If a man be without the virtues proper "
                          "to humanity, what has he to do with music?'",
                },
                {
                    "zh": "子貢欲去告朔之餼羊。子曰：「賜也！爾愛其羊，我愛其禮。」",
                    "en": "Tsze-kung wished to do away with the offering of a "
                          "sheep connected with the inauguration of the first "
                          "day of each month. The Master said, 'Ts'ze, you love "
                          "the sheep; I love the ceremony.'",
                },
                {
                    "zh": "子曰：「居上不寬，為禮不敬，臨喪不哀，吾何以觀之哉？」",
                    "en": "The Master said, 'High station filled without "
                          "indulgent generosity; ceremonies performed without "
                          "reverence; mourning conducted without sorrow; — "
                          "wherewith should I contemplate such ways?'",
                },
            ],
        },
    ],
}

BOOKS = {"analects": ANALECTS}


# --------------------------------------------------------------------------
# Rendering
# --------------------------------------------------------------------------

def chapter_html(part):
    """One chapter's passages. Shared by the DOCX, EPUB and PDF paths."""
    blocks = [
        f"<section class='passage'>"
        f"<p class='zh'>{section['zh']}</p>"
        f"<p class='en'>{section['en']}</p>"
        f"</section>"
        for section in part["sections"]
    ]
    return f"<h2>{part['title']}</h2>{''.join(blocks)}"


def book_html(book):
    """The whole book in one document.

    Books stopped being split into parts on 2026-08-14. What were
    separate downloadable units are now chapters inside a single work —
    which is how the book was written, and what a reader expects to
    receive.
    """
    chapters = "".join(chapter_html(part) for part in book["parts"])
    return f"""
    <article>
      <h1>{book['title']}</h1>
      <p class="byline">{book['author']} · {book['translator']}</p>
      {chapters}
      <p class="colophon"><em>{book['colophon']}</em></p>
    </article>
    """


def write_docx(book, path):
    doc = Document()

    title = doc.add_heading(book["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    byline = doc.add_paragraph(f"{book['author']} · {book['translator']}")
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    byline.runs[0].italic = True

    for index, part in enumerate(book["parts"]):
        if index:
            doc.add_page_break()
        doc.add_heading(part["title"], level=1)
        for section in part["sections"]:
            zh = doc.add_paragraph(section["zh"])
            zh.runs[0].font.size = Pt(13)
            doc.add_paragraph(section["en"])
            doc.add_paragraph()

    colophon = doc.add_paragraph(book["colophon"])
    colophon.runs[0].italic = True
    colophon.runs[0].font.size = Pt(9)

    doc.save(path)


def write_epub(book, path):
    epub_book = epub.EpubBook()
    epub_book.set_identifier(f"noblesee-{book['slug']}")
    epub_book.set_title(book["title"])
    epub_book.set_language("en")
    epub_book.add_author(f"{book['author']} (trans. {book['translator']})")

    style = epub.EpubItem(
        uid="style",
        file_name="style/noblesee.css",
        media_type="text/css",
        content=(
            "body{font-family:Georgia,serif;line-height:1.6;margin:1em;}"
            "h1{font-size:1.6em;} h2{font-size:1.2em;color:#555;}"
            ".zh{font-size:1.15em;} .byline{font-style:italic;color:#666;}"
            ".passage{margin-bottom:1.5em;}"
            ".colophon{font-size:0.8em;color:#888;margin-top:2em;}"
        ),
    )
    epub_book.add_item(style)

    chapters = []
    for index, part in enumerate(book["parts"], start=1):
        chapter = epub.EpubHtml(
            title=part["title"], file_name=f"chapter-{index}.xhtml", lang="en"
        )
        # The title block only on the first chapter; after that the book
        # is already open and repeating it reads as a new book each time.
        opening = (
            f"<h1>{book['title']}</h1>"
            f"<p class='byline'>{book['author']} · {book['translator']}</p>"
            if index == 1
            else ""
        )
        closing = (
            f"<p class='colophon'><em>{book['colophon']}</em></p>"
            if index == len(book["parts"])
            else ""
        )
        chapter.content = (
            f"<html><head><title>{part['title']}</title></head>"
            f"<body><article>{opening}{chapter_html(part)}{closing}</article></body></html>"
        )
        chapter.add_item(style)
        epub_book.add_item(chapter)
        chapters.append(chapter)

    # A real table of contents now that a book has more than one chapter.
    epub_book.toc = tuple(
        epub.Link(c.file_name, p["title"], p["slug"])
        for c, p in zip(chapters, book["parts"])
    )
    epub_book.add_item(epub.EpubNcx())
    epub_book.add_item(epub.EpubNav())

    # The text is the first page, not the table of contents. The nav
    # document stays in the manifest, which is what EPUB 3 requires and
    # what feeds the reader's chapter display; it just is not where the
    # book opens.
    epub_book.spine = chapters

    epub.write_epub(path, epub_book, {})


def pdf_css(base_pt):
    return f"""
    @page {{ size: A5; margin: 2cm; }}
    body {{ font-family: Georgia, 'Droid Sans Fallback', serif;
             font-size: {base_pt}pt; line-height: 1.6; color: #1a1a1a; }}
    h1 {{ font-size: {base_pt + 8}pt; text-align: center; }}
    h2 {{ font-size: {base_pt + 3}pt; color: #444; page-break-before: always; }}
    h2:first-of-type {{ page-break-before: avoid; }}
    .byline {{ text-align: center; font-style: italic; color: #555;
                margin-bottom: 1.5em; }}
    .zh {{ font-size: {base_pt + 1}pt; margin-bottom: .4em; }}
    .en {{ margin-top: 0; }}
    .passage {{ margin-bottom: 1.4em; }}
    .colophon {{ font-size: {max(base_pt - 3, 8)}pt; color: #777;
                  margin-top: 2em; }}
    """


def write_pdf(book, path, base_pt):
    html = (
        f"<html><head><style>{pdf_css(base_pt)}</style></head>"
        f"<body>{book_html(book)}</body></html>"
    )
    HTML(string=html).write_pdf(path)


def page_count(book):
    """Pages at the standard size — what the credit price is derived from.

    A DOCX carries no reliable page count (pagination is a rendering
    decision, and python-docx writes no `<Pages>` property), so the
    standard PDF rendered from the same content is the honest proxy.
    """
    html = (
        f"<html><head><style>{pdf_css(PDF_VARIANTS['pdf-standard'])}</style></head>"
        f"<body>{book_html(book)}</body></html>"
    )
    return len(HTML(string=html).render().pages)


def write_cover(book, path):
    img = Image.new("RGB", (800, 1200), color=(28, 26, 23))
    draw = ImageDraw.Draw(img)
    try:
        font_cjk = ImageFont.truetype(CJK_FONT, 72)
        font_small = ImageFont.truetype(LATIN_FONT, 32)
    except OSError:
        font_cjk = font_small = ImageFont.load_default()

    def center(y, text, font, fill):
        bbox = draw.textbbox((0, 0), text, font=font)
        draw.text(((800 - (bbox[2] - bbox[0])) / 2, y), text, font=font, fill=fill)

    center(420, book["cover_zh"], font_cjk, (212, 175, 105))
    center(520, book["cover_en"], font_small, (230, 225, 215))
    center(1080, "NobleSee", font_small, (120, 116, 108))
    img.save(path, quality=90)


# --------------------------------------------------------------------------

def generate(book, force=False):
    book_dir = os.path.join(SEED_ROOT, book["slug"], "whole")
    os.makedirs(book_dir, exist_ok=True)

    def emit(path, fn, *args):
        if os.path.exists(path) and not force:
            print(f"  skip (exists): {os.path.relpath(path, REPO_ROOT)}")
            return
        fn(path, *args)
        print(f"  wrote: {os.path.relpath(path, REPO_ROOT)}")

    print(f"{book['title']}")
    emit(os.path.join(SEED_ROOT, book["slug"], "cover.jpg"),
         lambda p: write_cover(book, p))
    emit(os.path.join(book_dir, "master.docx"), lambda p: write_docx(book, p))
    emit(os.path.join(book_dir, "book.epub"), lambda p: write_epub(book, p))
    for variant, size in PDF_VARIANTS.items():
        emit(os.path.join(book_dir, f"{variant}.pdf"),
             lambda p, s=size: write_pdf(book, p, s))

    # Printed so the value can be copied into the seed, which is where
    # the price comes from.
    print(f"  pages (standard): {page_count(book)}")


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--book", choices=sorted(BOOKS), help="only this book")
    parser.add_argument("--force", action="store_true", help="overwrite existing")
    args = parser.parse_args()

    selected = [BOOKS[args.book]] if args.book else list(BOOKS.values())
    for book in selected:
        generate(book, force=args.force)
    return 0


if __name__ == "__main__":
    sys.exit(main())
